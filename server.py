import os
from flask import Flask, render_template, request, send_file
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.enum.section import WD_ORIENT
from datetime import datetime
from flask_cors import CORS

app = Flask(__name__)
CORS(app)
app.config['JSON_SORT_KEYS'] = False
app.config['JSON_AS_ASCII'] = False

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/submit', methods=['POST'])
def submit_data():
    print('true')
    json_data = request.json

    # Создание нового документа
    doc = Document()

    # Изменение ориентации страницы на альбомную
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height

    # Установка полей страницы
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    # Заголовок
    group = json_data[0]['group']
    for entry in json_data:
        if group != entry['group']: group = False

    if group:
        heading = doc.add_heading(f'Группа {group}')
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        heading_run = heading.runs[0]
        heading_run.font.name = 'Times New Roman'
        heading_run.font.color.rgb = RGBColor(0, 0, 0)
        heading_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        heading_run.font.size = Pt(18)

    # Создание таблицы
    if group: cols = 7
    else: cols = 8

    table = doc.add_table(rows=1, cols=cols)
    table.style = 'Table Grid'

    # Добавление заголовков
    hdr_cells = table.rows[0].cells
    headers = [
        'ФИО уч-ся (моб. телефон)', 
        'Дата рождения', 
        'Паспортные данные', 
        'ФИО (отец), где и кем работает, контактный телефон', 
        'ФИО (мать), где и кем работает, контактный телефон', 
        'Дом. телефон', 
        'Индекс, домашний адрес'
    ]
    if not group: headers.append('Группа')
    
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        run.font.size = Pt(12)
        run.font.bold = True

    # Добавление данных
    for i, entry in enumerate(json_data, start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = f"{i}. {entry.get('fio_student', '')}\n{entry.get('additional', '')}"
        row_cells[1].text = entry.get('birth_date', '')
        row_cells[2].text = entry.get('passport_data', '')
        row_cells[3].text = f"{entry.get('fatherData', '')}"
        row_cells[4].text = f"{entry.get('motherData', '')}"
        row_cells[5].text = entry.get('home_phone', '')
        row_cells[6].text = f"{entry.get('address', '')}"
        if not group: row_cells[7].text = f"{entry.get('group', '')}"

        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.name = 'Times New Roman'
                    font.size = Pt(12)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # Сохранение документа
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f'data_{timestamp}.docx'
    doc.save(filename)

    return send_file(filename, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
